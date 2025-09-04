import os
import pandas as pd
import subprocess
import platform
import tempfile
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
    CSVLoader,
)
from langchain.schema import Document
from docx import Document as DocxDocument

try:
    import win32com.client  # Windows 用
except ImportError:
    win32com = None


def load_doc_file(file_path: str):
    """
    跨平台加载 .doc 文件：
    1. Windows：用 Word 转换成 .docx，再解析
    2. Linux/Mac：用 LibreOffice 转换成 .docx，再解析
    3. 如果都失败，尝试 antiword
    """
    text = ""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"❌ 文件不存在: {file_path}")

    tmp_dir = tempfile.mkdtemp()
    converted_docx = os.path.join(tmp_dir, "converted.docx")

    try:
        system_name = platform.system()

        # Windows 下用 Word 转换
        if system_name == "Windows" and win32com:
            try:
                word = win32com.client.Dispatch("Word.Application")
                doc = word.Documents.Open(file_path)
                doc.SaveAs(converted_docx, FileFormat=16)  # 16 = docx
                doc.Close()
                word.Quit()
                text = _parse_docx(converted_docx)
                return [Document(page_content=text, metadata={"source": file_path})]
            except Exception as e:
                print(f"⚠️ Word 转换失败: {e}")

        # Linux / Mac 下用 LibreOffice 转换
        if system_name in ["Linux", "Darwin"]:
            try:
                subprocess.run(
                    ["soffice", "--headless", "--convert-to", "docx", "--outdir", tmp_dir, file_path],
                    check=True
                )
                text = _parse_docx(converted_docx)
                return [Document(page_content=text, metadata={"source": file_path})]
            except Exception as e:
                print(f"⚠️ LibreOffice 转换失败: {e}")

        # 兜底方案：antiword
        try:
            text = subprocess.check_output(["antiword", file_path]).decode("utf-8", errors="ignore")
            return [Document(page_content=text, metadata={"source": file_path})]
        except Exception as e:
            print(f"❌ antiword 解析失败: {e}")
            return []

    finally:
        if os.path.exists(converted_docx):
            os.remove(converted_docx)


def _parse_docx(file_path: str) -> str:
    """解析 .docx 文件"""
    doc = DocxDocument(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])


def load_documents_from_folder(folder_path: str, encoding: str = "utf-8"):
    """批量加载 txt / pdf / docx / doc / csv / xlsx 文件"""
    all_docs = []
    supported_exts = [".txt", ".pdf", ".docx", ".doc", ".csv", ".xlsx"]

    for root, _, files in os.walk(folder_path):
        for file in files:
            ext = os.path.splitext(file)[-1].lower()
            file_path = os.path.join(root, file)

            if ext not in supported_exts:
                print(f"⚠️ 跳过不支持的文件: {file_path}")
                continue

            try:
                if ext == ".txt":
                    loader = TextLoader(file_path, encoding=encoding)
                    docs = loader.load()

                elif ext == ".pdf":
                    loader = PyPDFLoader(file_path)
                    docs = loader.load()

                elif ext == ".docx":
                    loader = UnstructuredWordDocumentLoader(file_path)
                    docs = loader.load()

                elif ext == ".doc":
                    docs = load_doc_file(file_path)

                elif ext == ".csv":
                    loader = CSVLoader(file_path, encoding=encoding)
                    docs = loader.load()

                elif ext == ".xlsx":
                    df = pd.read_excel(file_path)
                    docs = []
                    for _, row in df.iterrows():
                        text = " | ".join([f"{col}: {row[col]}" for col in df.columns])
                        docs.append(Document(page_content=text, metadata={"source": file_path}))
                else:
                    docs = []

                all_docs.extend(docs)
                print(f"✅ 成功加载: {file_path} ({len(docs)} 个切片)")

            except Exception as e:
                print(f"❌ 加载失败: {file_path}, 错误: {e}")

    return all_docs



if __name__ == "__main__":
    folder = "../docs"  # 替换成你的文档目录
    documents = load_documents_from_folder(folder)

    print(f"\n📊 总共加载文档切片数: {len(documents)}")
    print("示例:", documents[0] if documents else "没有文档")
